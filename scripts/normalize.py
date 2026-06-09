"""
讀取 sync_meta.json，解析直企 xlsx，寫入 Supabase，生成 P1+P2 截圖上傳 Storage。

用法：
  python3 scripts/normalize.py                      # 讀 data/raw/sync_meta.json
  python3 scripts/normalize.py path/to/sync_meta.json

環境變數：
  SUPABASE_URL          — Supabase project URL
  SUPABASE_SERVICE_KEY  — service_role key
"""

import hashlib
import io
import json
import math
import os
import re
import subprocess
import sys
import tempfile
import xml.etree.ElementTree as ET
from pathlib import Path

import fitz
import openpyxl
from openpyxl.utils import get_column_letter
from openpyxl.worksheet.page import PageMargins
from openpyxl.worksheet.properties import PageSetupProperties
from PIL import Image, ImageOps
from dotenv import load_dotenv
from supabase import create_client, Client

load_dotenv()
load_dotenv(os.environ.get("ENV_FILE", ".env.local"), override=True)

SUPABASE_URL = os.environ["SUPABASE_URL"]
SUPABASE_KEY = os.environ["SUPABASE_SERVICE_KEY"]
supabase: Client = create_client(SUPABASE_URL, SUPABASE_KEY)

META_PATH       = Path(__file__).parent.parent / "data" / "raw" / "sync_meta.json"
STATIC_PREVIEWS = Path(__file__).parent.parent / "app" / "static" / "previews"

COUNTY_NORMALIZE = {
    "臺北市": "台北", "台北市": "台北",
    "新北市": "新北",
    "基隆市": "基隆",
    "宜蘭縣": "宜蘭",
    "桃園市": "桃園",
    "新竹市": "新竹", "新竹縣": "新竹",
    "苗栗縣": "苗栗",
    "臺中市": "台中", "台中市": "台中",
    "花蓮縣": "花蓮",
    "彰化縣": "彰化",
    "南投縣": "南投",
    "雲林縣": "雲林",
    "嘉義市": "嘉義", "嘉義縣": "嘉義",
    "臺南市": "台南", "台南市": "台南",
    "高雄市": "高雄",
    "屏東縣": "屏東",
    "臺東縣": "台東", "台東縣": "台東",
}

CONTENT_TYPE_MAP = {
    ".pdf":  "application/pdf",
    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ".jpg":  "image/jpeg",
    ".jpeg": "image/jpeg",
    ".png":  "image/png",
    ".gpx":  "application/gpx+xml",
    ".kml":  "application/vnd.google-earth.kml+xml",
    ".txt":  "text/plain; charset=utf-8",
    ".md":   "text/markdown; charset=utf-8",
}


# ── 工具函式 ─────────────────────────────────────────────────

def roc_to_iso(text: str) -> str | None:
    m = re.search(r"(\d+)\s*年\s*(\d+)\s*月\s*(\d+)\s*日", str(text))
    if not m:
        return None
    y = int(m.group(1)) + 1911
    return f"{y:04d}-{int(m.group(2)):02d}-{int(m.group(3)):02d}"


def extract_county_town(location: str) -> tuple[str | None, str | None]:
    county = None
    for official, display in COUNTY_NORMALIZE.items():
        if official in location:
            county = display
            break
    town = None
    m = re.search(r"[縣市](.{1,6}?[鄉鎮市區])", location)
    if m:
        town = m.group(1)
    return county, town


def cell_value(ws, coord: str) -> str:
    return str(ws[coord].value or "").strip()


def storage_safe_name(filename: str) -> str:
    suffix = Path(filename).suffix.lower()
    try:
        filename.encode("ascii")
        safe = re.sub(r"[^\w\-.]", "_", Path(filename).stem, flags=re.ASCII)
        safe = re.sub(r"_{2,}", "_", safe).strip("_") or "file"
        return f"{safe}{suffix}"
    except UnicodeEncodeError:
        name_hash = hashlib.sha1(filename.encode()).hexdigest()[:12]
        return f"{name_hash}{suffix}"


_GPX_NS = "http://www.topografix.com/GPX/1/1"


def _rdp_keep(pts: list[tuple[float, float]], eps: float) -> set[int]:
    n = len(pts)
    if n < 3:
        return set(range(n))
    keep = bytearray(n)
    keep[0] = keep[n - 1] = 1
    stack = [(0, n - 1)]
    while stack:
        s, e = stack.pop()
        dx, dy = pts[e][0] - pts[s][0], pts[e][1] - pts[s][1]
        ls = dx * dx + dy * dy
        max_d, idx = 0.0, s
        for i in range(s + 1, e):
            t = max(0.0, min(1.0, ((pts[i][0]-pts[s][0])*dx + (pts[i][1]-pts[s][1])*dy) / ls)) if ls else 0.0
            d = math.hypot(pts[i][0]-pts[s][0]-t*dx, pts[i][1]-pts[s][1]-t*dy)
            if d > max_d:
                max_d, idx = d, i
        if max_d > eps:
            keep[idx] = 1
            stack.extend([(s, idx), (idx, e)])
    return {i for i, v in enumerate(keep) if v}


def simplify_gpx(data: bytes, epsilon: float = 0.0001) -> tuple[bytes, int, int]:
    """RDP simplification + 5-decimal rounding on GPX trkpt data.
    Returns (simplified_bytes, points_before, points_after).
    """
    namespaces = {p: u for _, (p, u) in ET.iterparse(io.BytesIO(data), events=["start-ns"])}
    for prefix, uri in namespaces.items():
        ET.register_namespace(prefix, uri)
    ET.register_namespace("", _GPX_NS)

    root = ET.fromstring(data)
    total_before = total_after = 0

    for trkseg in root.findall(f".//{{{_GPX_NS}}}trkseg"):
        trkpts = trkseg.findall(f"{{{_GPX_NS}}}trkpt")
        if not trkpts:
            continue
        pts = [(float(t.get("lat")), float(t.get("lon"))) for t in trkpts]
        kept = _rdp_keep(pts, epsilon)
        total_before += len(trkpts)
        total_after += len(kept)
        for i, trkpt in enumerate(trkpts):
            if i in kept:
                trkpt.set("lat", str(round(float(trkpt.get("lat")), 5)))
                trkpt.set("lon", str(round(float(trkpt.get("lon")), 5)))
            else:
                trkseg.remove(trkpt)

    xml_str = '<?xml version="1.0" encoding="UTF-8"?>\n' + ET.tostring(root, encoding="unicode")
    return xml_str.encode("utf-8"), total_before, total_after


def storage_upload(bucket: str, path: str, source: "Path | bytes", content_type: str):
    data = source if isinstance(source, bytes) else source.read_bytes()
    try:
        supabase.storage.from_(bucket).upload(
            path, data,
            file_options={"content-type": content_type, "upsert": "true"},
        )
    except Exception as e:
        print(f"    ⚠ 上傳 {bucket}/{path} 失敗：{e}")


# ── 截圖 ─────────────────────────────────────────────────────

def find_content_bounds(ws) -> tuple[int, int]:
    max_row, max_col = 1, 1
    for row in ws.iter_rows():
        for cell in row:
            if cell.value is not None:
                max_row = max(max_row, cell.row)
                max_col = max(max_col, cell.column)
    return max_row, max_col


def capture_sheet(xlsx_path: Path, sheet_name: str, cell_range: str | None, output_path: Path):
    """
    將指定 sheet 轉成 PNG。
    cell_range=None：自動偵測內容邊界（P1）。
    cell_range="B2:O12"：固定範圍（P2）。
    """
    with tempfile.TemporaryDirectory() as _tmp:
        tmp = Path(_tmp)
        wb = openpyxl.load_workbook(xlsx_path, data_only=True)
        ws = wb[sheet_name]

        if cell_range is None:
            max_row, max_col = find_content_bounds(ws)
            cell_range = f"A1:{get_column_letter(max_col)}{max_row}"

        ws.print_area = cell_range
        if ws.sheet_properties.pageSetUpPr is None:
            ws.sheet_properties.pageSetUpPr = PageSetupProperties(fitToPage=True)
        else:
            ws.sheet_properties.pageSetUpPr.fitToPage = True
        ws.page_setup.fitToWidth = 1
        ws.page_setup.fitToHeight = 1
        ws.page_setup.scale = None
        ws.page_margins = PageMargins(left=0.2, right=0.2, top=0.2, bottom=0.2, header=0, footer=0)

        for name in list(wb.sheetnames):
            if name != sheet_name:
                del wb[name]

        tmp_xlsx = tmp / "sheet.xlsx"
        wb.save(tmp_xlsx)

        result = subprocess.run(
            ["libreoffice", "--headless", "--convert-to", "pdf",
             str(tmp_xlsx), "--outdir", str(tmp)],
            capture_output=True, timeout=60,
        )
        if result.returncode != 0:
            print(f"    ⚠ LibreOffice 失敗：{result.stderr.decode()[:200]}")
            return

        pdf_path = tmp / "sheet.pdf"
        if not pdf_path.exists():
            print(f"    ⚠ PDF 未生成")
            return

        doc = fitz.open(pdf_path)
        pix = doc[0].get_pixmap(matrix=fitz.Matrix(2.0, 2.0))
        output_path.parent.mkdir(parents=True, exist_ok=True)
        pix.save(str(output_path))


def trim_whitespace(img: Image.Image, padding: int = 15) -> Image.Image:
    gray = img.convert("L")
    bbox = ImageOps.invert(gray).getbbox()
    if not bbox:
        return img
    l, t, r, b = bbox
    return img.crop((max(0, l - padding), max(0, t - padding),
                     min(img.width, r + padding), min(img.height, b + padding)))


def build_preview(p1_path: Path, p2_path: Path, output_path: Path):
    A4_W, A4_H, GAP = 1240, 1754, 16
    imgs = []
    for p in (p1_path, p2_path):
        if p.exists():
            imgs.append(trim_whitespace(Image.open(p)))
    if not imgs:
        return

    resized = [img.resize((A4_W, round(img.height * A4_W / img.width)), Image.LANCZOS)
               for img in imgs]
    total_h = sum(i.height for i in resized) + GAP * (len(resized) - 1)
    if total_h > A4_H:
        scale = A4_H / total_h
        w = round(A4_W * scale)
        resized = [img.resize((w, round(img.height * w / img.width)), Image.LANCZOS)
                   for img in resized]
        total_h = sum(i.height for i in resized) + GAP * (len(resized) - 1)

    canvas = Image.new("RGB", (max(i.width for i in resized), total_h), "white")
    y = 0
    for img in resized:
        canvas.paste(img, (0, y))
        y += img.height + GAP
    output_path.parent.mkdir(parents=True, exist_ok=True)
    canvas.save(str(output_path))


# ── 直企解析 ─────────────────────────────────────────────────

P1_NAMES = ["直企P1(列印)", "直企列印 P1"]
P2_NAMES = ["直企P2(列印)", "直企列印 P2"]


def find_sheet(wb, candidates: list[str]) -> str | None:
    for name in candidates:
        if name in wb.sheetnames:
            return name
    return None


def parse_p1(ws) -> dict:
    name       = cell_value(ws, "D2")
    date_start = roc_to_iso(cell_value(ws, "C3"))
    date_end   = roc_to_iso(cell_value(ws, "C4"))
    entry_loc  = cell_value(ws, "F3")
    exit_loc   = cell_value(ws, "F4")
    leader_display = cell_value(ws, "C17") or None
    entry_county, entry_town = extract_county_town(entry_loc)
    exit_county,  exit_town  = extract_county_town(exit_loc)
    return {
        "name":                name or None,
        "date_start":          date_start,
        "date_end":            date_end,
        "region_entry_county": entry_county,
        "region_entry_town":   entry_town,
        "region_exit_county":  exit_county,
        "region_exit_town":    exit_town,
        "leader_display":      leader_display,
    }


# ── DB 操作 ──────────────────────────────────────────────────

def upsert_group(drive_folder_id: str, name: str) -> int:
    existing = supabase.table("expedition_groups").select("id").eq("drive_folder_id", drive_folder_id).execute()
    if existing.data:
        return existing.data[0]["id"]
    result = supabase.table("expedition_groups").insert({
        "name": name,
        "drive_folder_id": drive_folder_id,
    }).execute()
    return result.data[0]["id"]


def upsert_expedition(drive_folder_id: str, group_id: int, fields: dict) -> tuple[int, bool]:
    existing = supabase.table("expeditions").select("id").eq("drive_folder_id", drive_folder_id).execute()
    if existing.data:
        return existing.data[0]["id"], False
    result = supabase.table("expeditions").insert({
        "drive_folder_id": drive_folder_id,
        "group_id": group_id,
        **fields,
    }).execute()
    return result.data[0]["id"], True


def sync_file_record(table: str, expedition_id: int, drive_file_id: str,
                     filename: str, extra: dict):
    existing = supabase.table(table).select("id").eq("drive_file_id", drive_file_id).execute()
    if existing.data:
        supabase.table(table).update({"filename": filename, **extra}).eq("drive_file_id", drive_file_id).execute()
    else:
        supabase.table(table).insert({
            "expedition_id": expedition_id,
            "drive_file_id": drive_file_id,
            "filename": filename,
            **extra,
        }).execute()


def sync_counties(expedition_id: int, counties: set[str]):
    for county in counties:
        supabase.table("expedition_counties").upsert(
            {"expedition_id": expedition_id, "county": county},
            on_conflict="expedition_id,county",
        ).execute()


# ── 靜態檔案 ─────────────────────────────────────────────────

def process_gpx_files(expedition_id: int, file_entries: list[dict]):
    for f in file_entries:
        local = Path(f["local_path"])
        if not local.exists():
            continue
        safe = storage_safe_name(f["name"])
        storage_path = f"{expedition_id}/{safe}"
        content_type = CONTENT_TYPE_MAP.get(local.suffix.lower(), "application/octet-stream")
        if local.suffix.lower() == ".gpx":
            simplified, before, after = simplify_gpx(local.read_bytes())
            pct = (1 - after / before) * 100 if before else 0
            print(f"    → GPX RDP: {before:,} → {after:,} pts ({pct:.0f}% removed)")
            storage_upload("gpx", storage_path, simplified, content_type)
        else:
            storage_upload("gpx", storage_path, local, content_type)
        sync_file_record("gpx_files", expedition_id, f["drive_file_id"],
                         f["name"], {"file_path": storage_path})


def process_map_files(expedition_id: int, file_entries: list[dict]):
    for f in file_entries:
        local = Path(f["local_path"])
        if not local.exists():
            continue
        safe = storage_safe_name(f["name"])
        storage_path = f"{expedition_id}/{safe}"
        content_type = CONTENT_TYPE_MAP.get(local.suffix.lower(), "application/octet-stream")
        storage_upload("maps", storage_path, local, content_type)
        sync_file_record("map_files", expedition_id, f["drive_file_id"],
                         f["name"], {"file_path": storage_path})


def process_record_files(expedition_id: int, file_entries: list[dict]):
    for f in file_entries:
        local = Path(f["local_path"])
        if not local.exists():
            continue
        ext = local.suffix.lower()
        if ext == ".docx":
            from docx import Document
            doc = Document(local)
            parts = [p.text for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    row_text = "  ".join(c.text.strip() for c in row.cells if c.text.strip())
                    if row_text:
                        parts.append(row_text)
            content = "\n".join(parts)
        elif ext == ".pdf":
            doc = fitz.open(str(local))
            content = "\n".join(page.get_text() for page in doc)
            doc.close()
        else:
            content = local.read_text(encoding="utf-8", errors="replace")
        safe = storage_safe_name(f["name"])
        storage_path = f"{expedition_id}/{safe}"
        content_type = CONTENT_TYPE_MAP.get(ext, "application/octet-stream")
        storage_upload("records", storage_path, local, content_type)
        sync_file_record("records", expedition_id, f["drive_file_id"],
                         f["name"], {"content": content, "file_path": storage_path})


# ── 截圖與上傳 ───────────────────────────────────────────────

def generate_and_upload_preview(xlsx_path: Path, exp_id: int):
    STATIC_PREVIEWS.mkdir(parents=True, exist_ok=True)
    preview_local = STATIC_PREVIEWS / f"{exp_id}.png"

    wb = openpyxl.load_workbook(xlsx_path, data_only=True)
    p1_name = find_sheet(wb, P1_NAMES)
    p2_name = find_sheet(wb, P2_NAMES)

    with tempfile.TemporaryDirectory() as _tmp:
        tmp = Path(_tmp)
        p1_png = tmp / "p1.png"
        p2_png = tmp / "p2.png"

        if p1_name:
            print(f"    截圖 P1...", end=" ", flush=True)
            capture_sheet(xlsx_path, p1_name, None, p1_png)
            print("完成" if p1_png.exists() else "失敗")

        if p2_name:
            print(f"    截圖 P2...", end=" ", flush=True)
            capture_sheet(xlsx_path, p2_name, "B2:O12", p2_png)
            print("完成" if p2_png.exists() else "失敗")

        build_preview(p1_png, p2_png, preview_local)

    if preview_local.exists():
        storage_upload("previews", f"{exp_id}.png", preview_local, "image/png")
        supabase.table("expeditions").update({"preview_image": f"{exp_id}.png"}).eq("id", exp_id).execute()
        print(f"    預覽圖：{exp_id}.png")


# ── 主流程 ───────────────────────────────────────────────────

def process_expedition(entry: dict, group_id: int):
    xlsx_info = entry.get("xlsx")
    if not xlsx_info:
        print(f"  ⚠ 無直企資訊，跳過 {entry['name']}")
        return

    xlsx_path = Path(xlsx_info["local_path"])
    if not xlsx_path.exists():
        print(f"  ⚠ 找不到 xlsx：{xlsx_path}，跳過")
        return

    wb = openpyxl.load_workbook(xlsx_path, data_only=True)
    p1_name = find_sheet(wb, P1_NAMES)
    if not p1_name:
        print(f"  ⚠ 找不到 P1 sheet，跳過 {entry['name']}")
        return

    fields = parse_p1(wb[p1_name])
    fields["name"] = entry["name"]
    if not fields["date_start"]:
        print(f"  ⚠ 無法解析 date_start，跳過")
        return

    exp_id, is_new = upsert_expedition(entry["drive_folder_id"], group_id, fields)

    counties = {c for c in (fields["region_entry_county"], fields["region_exit_county"]) if c}
    if counties:
        sync_counties(exp_id, counties)

    process_gpx_files(exp_id, entry.get("gpx_files", []))
    process_map_files(exp_id, entry.get("map_files", []))
    process_record_files(exp_id, entry.get("record_files", []))

    if is_new:
        print(f"  ✓ 新增：{entry['name']}（id={exp_id}）")
        generate_and_upload_preview(xlsx_path, exp_id)
    else:
        print(f"  → 更新檔案：{entry['name']}（id={exp_id}）")


def update_last_synced_at(synced_at: str):
    supabase.table("sync_state").upsert(
        {"key": "last_synced_at", "value": synced_at},
        on_conflict="key",
    ).execute()


def _append_normalize_log(log_text: str):
    try:
        row = supabase.table("sync_logs").select("id, log_text").order("id", desc=True).limit(1).execute()
        if not row.data:
            return
        existing = row.data[0].get("log_text") or ""
        supabase.table("sync_logs").update({
            "log_text": existing + "\n[normalize]\n" + log_text,
        }).eq("id", row.data[0]["id"]).execute()
    except Exception as e:
        print(f"  ⚠ 無法更新 sync_logs：{e}", file=sys.stderr)


class _Tee:
    def __init__(self, real):
        self._real = real
        self._buf = io.StringIO()
    def write(self, s):
        self._real.write(s)
        self._buf.write(s)
    def flush(self):
        self._real.flush()
    def getvalue(self) -> str:
        return self._buf.getvalue()


def main():
    _tee = _Tee(sys.stdout)
    sys.stdout = _tee

    meta_path = Path(sys.argv[1]) if len(sys.argv) >= 2 else META_PATH
    if not meta_path.exists():
        sys.stdout = _tee._real
        print(f"sync_meta.json 不存在：{meta_path}", file=sys.stderr)
        sys.exit(1)

    meta = json.loads(meta_path.read_text(encoding="utf-8"))

    for solo in meta.get("solos", []):
        print(f"\n處理（solo）：{solo['name']}")
        try:
            group_id = upsert_group(solo["drive_folder_id"], solo["name"])
            process_expedition(solo, group_id)
        except Exception as e:
            print(f"  ✗ 錯誤：{e}")

    for group in meta.get("groups", []):
        print(f"\n處理（group）：{group['name']}")
        try:
            group_id = upsert_group(group["drive_folder_id"], group["name"])
            for team in group.get("teams", []):
                print(f"  隊伍：{team['name']}")
                process_expedition(team, group_id)
        except Exception as e:
            print(f"  ✗ 錯誤：{e}")

    update_last_synced_at(meta["synced_at"])
    print(f"\nnormalize complete, last_synced_at → {meta['synced_at']}")

    sys.stdout = _tee._real
    _append_normalize_log(_tee.getvalue())


if __name__ == "__main__":
    main()
