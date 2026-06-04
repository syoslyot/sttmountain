import os, sys, sqlite3
from pathlib import Path
import fitz
import docx as docxlib
sys.path.insert(0, str(Path(__file__).parent))
from normalize import capture_sheet_range, build_a4_preview

BASE       = Path(__file__).parent.parent
DB_PATH    = BASE / 'db/sttmount.db'
PREVIEWS   = BASE / 'app/static/previews'
RAW        = BASE / 'data/raw'
XLSX_DIR   = BASE / 'data/raw'

def get_conn():
    conn = sqlite3.connect(DB_PATH)
    conn.execute("PRAGMA foreign_keys = ON")
    return conn

def pdf_text(path):
    doc = fitz.open(str(path))
    return '\n'.join(page.get_text() for page in doc).strip()

def docx_text(path):
    doc = docxlib.Document(str(path))
    lines = [p.text for p in doc.paragraphs if p.text.strip()]
    for t in doc.tables:
        for row in t.rows:
            row_text = '\t'.join(c.text.strip() for c in row.cells)
            if row_text.strip():
                lines.append(row_text)
    return '\n'.join(lines)

def excel_preview(xlsx_path, out_png):
    import tempfile
    xlsx_path = Path(xlsx_path)
    with tempfile.TemporaryDirectory() as tmp:
        tmp = Path(tmp)
        p1, p2 = tmp / 'p1.png', tmp / 'p2.png'
        capture_sheet_range(xlsx_path, '直企P1(列印)', 'A2:G27', p1)
        capture_sheet_range(xlsx_path, '直企P2(列印)', 'B2:O11', p2)
        build_a4_preview([p for p in [p1, p2] if p.exists()], Path(out_png))
    return Path(out_png).exists()

def ins(conn, name, date_start, date_end, county, region, region_exit, leader_display,
        extra_counties, gpx_list, map_list, rec_list):
    cur = conn.execute(
        "INSERT INTO expeditions(name,date_start,date_end,county,region,region_exit,leader_display) VALUES (?,?,?,?,?,?,?)",
        (name, date_start, date_end, county, region, region_exit, leader_display)
    )
    eid = cur.lastrowid
    for c in [county] + (extra_counties or []):
        conn.execute("INSERT OR IGNORE INTO expedition_counties(expedition_id,county) VALUES (?,?)", (eid, c))
    for fname in gpx_list:
        conn.execute("INSERT INTO gpx_files(expedition_id,filename,file_path) VALUES (?,?,?)", (eid, fname, fname))
    for fname, ftype in map_list:
        conn.execute("INSERT INTO map_files(expedition_id,filename,file_path,file_type) VALUES (?,?,?,?)", (eid, fname, fname, ftype))
    for fname, content in rec_list:
        conn.execute("INSERT INTO records(expedition_id,filename,content) VALUES (?,?,?)", (eid, fname, content))
    return eid

def main():
    conn = get_conn()

    print('提取紀錄文字...')
    rec1_text = pdf_text(RAW / 'txt/260501-茶茶牙頓喝茶茶紀錄.pdf')
    rec2_text = docx_text(RAW / 'txt/白雪村紀錄＿林咨筌.docx')

    print('插入資料...')
    id1 = ins(conn,
        name='茶茶牙頓出西都驕溪',
        date_start='2026-04-30', date_end='2026-05-03',
        county='台東', region='臺東縣達仁鄉', region_exit='屏東縣獅子鄉',
        leader_display='程效賢',
        extra_counties=['屏東'],
        gpx_list=['茶茶牙頓喝茶茶航跡_福利熊.gpx'],
        map_list=[('A3茶茶牙頓地圖.pdf','pdf'), ('A4西都驕溪地圖.pdf','pdf')],
        rec_list=[('260501-茶茶牙頓喝茶茶紀錄.pdf', rec1_text)],
    )

    id2 = ins(conn,
        name='白雪村採樟腦（水山線）',
        date_start='2024-03-20', date_end='2024-03-22',
        county='南投', region='南投縣信義鄉', region_exit='嘉義縣阿里山鄉',
        leader_display='李思誼',
        extra_counties=['嘉義'],
        gpx_list=['阿里山：水山霞山_薰慈.GPX'],
        map_list=[('地圖_航機.jpg','image')],
        rec_list=[('白雪村紀錄＿林咨筌.docx', rec2_text)],
    )
    conn.commit()
    print(f'  茶茶牙頓 id={id1}，白雪村 id={id2}')

    print('產生預覽圖...')
    xlsx1 = XLSX_DIR / '茶茶牙頓出西都驕溪_all in one直企.xlsx'
    xlsx2 = XLSX_DIR / '20260321_白雪村採樟腦（水山線）_直企 .xlsx'

    for eid, xlsx in [(id1, xlsx1), (id2, xlsx2)]:
        out = PREVIEWS / f'{eid}.png'
        if excel_preview(xlsx, out):
            conn.execute("UPDATE expeditions SET preview_image=? WHERE id=?", (f'previews/{eid}.png', eid))
            print(f'  id={eid} 預覽圖 OK → {out.name}')
        else:
            print(f'  id={eid} 預覽圖失敗')

    conn.commit()
    conn.close()
    print('完成。')

if __name__ == '__main__':
    main()
